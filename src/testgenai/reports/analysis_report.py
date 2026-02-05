from __future__ import annotations

from typing import List

from docx import Document

from testgenai.models.requirement import Requirement


def write_requirements_report(requirements: List[Requirement], out_path: str) -> None:
    doc = Document()
    doc.add_heading("Requirements Analysis", level=1)
    for req in requirements:
        doc.add_paragraph(f"{req.req_id}: {req.title}")
        doc.add_paragraph(req.description)
    doc.save(out_path)
