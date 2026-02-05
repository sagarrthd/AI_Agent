from __future__ import annotations

from typing import List

from docx import Document

from testgenai.models.scenario import Scenario


def write_scenarios_doc(scenarios: List[Scenario], out_path: str) -> None:
    doc = Document()
    doc.add_heading("High-Level Scenarios", level=1)
    for sc in scenarios:
        doc.add_paragraph(f"{sc.scenario_id}: {sc.title}")
        doc.add_paragraph(sc.description)
    doc.save(out_path)
